import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from typing import Dict, List, Tuple
from sklearn.impute import SimpleImputer
import google.generativeai as genai
from io import BytesIO
from zipfile import ZipFile
from datetime import datetime
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
import tempfile
import os

class MissingValueHandler:
    def __init__(self, df: pd.DataFrame):
        self.df = df.copy()
        self.numeric_cols = self.df.select_dtypes(include=[np.number]).columns
        self.categorical_cols = self.df.select_dtypes(include=['object']).columns
        
    def get_missing_info(self) -> pd.DataFrame:
        """Get information about missing values"""
        missing_info = pd.DataFrame({
            'Total Missing': self.df.isnull().sum(),
            'Percent Missing': (self.df.isnull().sum() / len(self.df) * 100).round(2)
        })
        return missing_info
    
    def impute_values(self, method: str, columns: List[str] = None) -> pd.DataFrame:
        """
        Impute missing values using specified method
        Methods: mean, median, mode, constant, ffill, bfill
        """
        df_imputed = self.df.copy()
        
        if columns is None:
            numeric_cols = self.numeric_cols
            categorical_cols = self.categorical_cols
        else:
            numeric_cols = [col for col in columns if col in self.numeric_cols]
            categorical_cols = [col for col in columns if col in self.categorical_cols]
        
        if method in ['mean', 'median']:
            if numeric_cols.empty:
                return df_imputed, "No numeric columns to impute"
                
            imputer = SimpleImputer(
                strategy=method,
                missing_values=np.nan
            )
            df_imputed[numeric_cols] = imputer.fit_transform(df_imputed[numeric_cols])
            
        elif method == 'mode':
            imputer = SimpleImputer(
                strategy='most_frequent',
                missing_values=np.nan
            )
            if not numeric_cols.empty:
                df_imputed[numeric_cols] = imputer.fit_transform(df_imputed[numeric_cols])
            if not categorical_cols.empty:
                df_imputed[categorical_cols] = imputer.fit_transform(df_imputed[categorical_cols])
                
        elif method == 'constant':
            for col in numeric_cols:
                df_imputed[col] = df_imputed[col].fillna(0)
            for col in categorical_cols:
                df_imputed[col] = df_imputed[col].fillna('MISSING')
                
        elif method == 'ffill':
            df_imputed = df_imputed.fillna(method='ffill')
            
        elif method == 'bfill':
            df_imputed = df_imputed.fillna(method='bfill')
            
        return df_imputed, f"Missing values imputed using {method} method"

class InsightsGenerator:
    def __init__(self, api_key: str):
        genai.configure(api_key="AIzaSyDjCfjNeoJjk0RUxsLAHcbB8pUtdfRBTgE")
        self.model = genai.GenerativeModel('gemini-pro')
        
    def generate_insights(self, df: pd.DataFrame) -> str:
        """Generate comprehensive insights about the dataset"""
        # First get statistical insights
        stats_insights = self._generate_statistical_insights(df)
        
        # Then get AI-powered analytical insights
        try:
            analytical_insights = self._generate_analytical_insights(stats_insights)
            return f"{stats_insights}\n\nAnalytical Insights:\n{analytical_insights}"
        except Exception as e:
            # Fallback to statistical insights if AI analysis fails
            return stats_insights
    
    def _generate_statistical_insights(self, df: pd.DataFrame) -> str:
        """Generate statistical insights about the dataset"""
        insights = []
        
        # Dataset overview
        insights.append(f"Dataset Shape: {df.shape[0]} rows and {df.shape[1]} columns")
        
        # Missing values
        missing_counts = df.isnull().sum()
        if missing_counts.any():
            cols_with_missing = missing_counts[missing_counts > 0]
            insights.append("\nMissing Values:")
            for col, count in cols_with_missing.items():
                percentage = (count / len(df) * 100)
                if not np.isnan(percentage):
                    percentage = round(percentage, 2)
                insights.append(f"- {col}: {count} values ({percentage}%)")
        
        # Numeric columns summary
        numeric_cols = df.select_dtypes(include=[np.number]).columns
        if len(numeric_cols) > 0:
            insights.append("\nNumeric Columns Summary:")
            for col in numeric_cols:
                mean = df[col].mean()
                median = df[col].median()
                std = df[col].std()
                if not np.isnan(mean):
                    mean = round(mean, 2)
                if not np.isnan(median):
                    median = round(median, 2)
                if not np.isnan(std):
                    std = round(std, 2)
                insights.append(f"- {col}: mean={mean}, median={median}, std={std}")
        
        # Categorical columns summary
        categorical_cols = df.select_dtypes(include=['object']).columns
        if len(categorical_cols) > 0:
            insights.append("\nCategorical Columns Summary:")
            for col in categorical_cols:
                unique_count = df[col].nunique()
                mode_values = df[col].mode()
                top_value = mode_values.iloc[0] if not mode_values.empty else "N/A"
                insights.append(f"- {col}: {unique_count} unique values, most common: {top_value}")
        
        return "\n".join(insights)
    
    def _generate_analytical_insights(self, stats_text: str) -> str:
        """Generate analytical insights using AI"""
        prompt = f"""
        As a data analyst, provide clear and concise insights about this dataset. 
        Focus on key trends, patterns, and notable observations.
        
        Statistical Information:
        {stats_text}
        
        Provide 3-4 key insights in natural language, focusing on:
        1. Data quality and completeness
        2. Distribution patterns and potential outliers
        3. Relationships between variables (if applicable)
        4. Any notable patterns or anomalies
        
        Keep the response concise and data-driven.
        """
        
        response = self.model.generate_content(prompt)
        return response.text

    def suggest_visualizations(self, df: pd.DataFrame) -> str:
        """Suggest appropriate visualizations based on the dataset structure"""
        # Prepare dataset information
        numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
        categorical_cols = df.select_dtypes(include=['object']).columns.tolist()
        
        prompt = f"""
        As a data visualization expert, suggest appropriate charts and graphs for this dataset.
        
        Dataset Information:
        - Numeric columns: {numeric_cols}
        - Categorical columns: {categorical_cols}
        - Total rows: {len(df)}
        
        For each suggestion:
        1. Specify the type of visualization
        2. Which columns to use
        3. What insights it could reveal
        
        Focus on:
        - Distribution analysis
        - Relationship analysis
        - Comparison analysis
        - Time series analysis (if applicable)
        
        Keep suggestions practical and data-driven. Limit to 4-5 most relevant visualizations.
        """
        
        try:
            response = self.model.generate_content(prompt)
            return response.text
        except Exception as e:
            # Fallback to basic suggestions if AI fails
            return self._generate_basic_visualization_suggestions(df)
    
    def _generate_basic_visualization_suggestions(self, df: pd.DataFrame) -> str:
        """Generate basic visualization suggestions without AI"""
        suggestions = []
        numeric_cols = df.select_dtypes(include=[np.number]).columns
        categorical_cols = df.select_dtypes(include=['object']).columns
        
        # Distribution analysis
        if len(numeric_cols) > 0:
            suggestions.append("\nFor Distribution Analysis:")
            suggestions.append(f"- Histograms for numeric columns: {', '.join(numeric_cols)}")
            suggestions.append(f"- Box plots for numeric columns to identify outliers")
        
        # Relationship analysis
        if len(numeric_cols) >= 2:
            suggestions.append("\nFor Relationship Analysis:")
            suggestions.append("- Correlation heatmap for numeric columns")
            suggestions.append("- Scatter plots for pairs of numeric columns")
        
        # Categorical analysis
        if len(categorical_cols) > 0:
            suggestions.append("\nFor Categorical Analysis:")
            suggestions.append(f"- Bar charts for categorical columns: {', '.join(categorical_cols)}")
            
            if len(numeric_cols) > 0:
                suggestions.append("- Box plots or violin plots to show distribution across categories")
        
        return "\n".join(suggestions)

    def suggest_imputation_strategy(self, df: pd.DataFrame) -> str:
        """Suggest the best imputation strategy for missing values in the dataset"""
        # Prepare dataset information
        missing_info = pd.DataFrame({
            'Total Missing': df.isnull().sum(),
            'Percent Missing': (df.isnull().sum() / len(df) * 100).round(2)
        })
        
        # Only analyze columns with missing values
        cols_with_missing = missing_info[missing_info['Total Missing'] > 0]
        
        if cols_with_missing.empty:
            return "No missing values found in the dataset."
        
        # Prepare column-specific information
        column_info = []
        for col in cols_with_missing.index:
            col_data = df[col].dropna()
            data_type = df[col].dtype
            unique_count = df[col].nunique()
            
            # Calculate additional statistics for numeric columns
            if np.issubdtype(data_type, np.number):
                skewness = col_data.skew() if len(col_data) > 0 else None
                has_outliers = any(np.abs((col_data - col_data.mean()) / col_data.std()) > 3) if len(col_data) > 0 else None
                
                # Fix the formatting issue
                skewness_str = f"{skewness:.2f}" if skewness is not None else "N/A"
                
                column_info.append(f"""
                Column: {col}
                - Data type: {data_type}
                - Missing values: {cols_with_missing.loc[col, 'Total Missing']} ({cols_with_missing.loc[col, 'Percent Missing']}%)
                - Unique values: {unique_count}
                - Skewness: {skewness_str}
                - Contains outliers: {has_outliers if has_outliers is not None else 'N/A'}
                """)
            else:
                column_info.append(f"""
                Column: {col}
                - Data type: {data_type}
                - Missing values: {cols_with_missing.loc[col, 'Total Missing']} ({cols_with_missing.loc[col, 'Percent Missing']}%)
                - Unique values: {unique_count}
                """)
        
        prompt = f"""
        As a data scientist, suggest the best imputation strategy for missing values in this dataset.
        
        Dataset Information:
        {'\n'.join(column_info)}
        
        For each column with missing values, recommend:
        1. The best imputation method (mean, median, mode, forward fill, backward fill, or constant)
        2. Justify why this method is most appropriate
        3. Potential risks or considerations for the chosen method
        
        Consider these factors in your recommendation:
        - Data type of each column
        - Distribution and skewness (for numeric data)
        - Presence of outliers
        - Percentage of missing values
        - Nature of the data (if apparent from column names)
        
        Provide a clear, concise recommendation for each column.
        """
        
        try:
            response = self.model.generate_content(prompt)
            return response.text
        except Exception as e:
            # Fallback to basic suggestions if AI fails
            return self._generate_basic_imputation_suggestions(df)
    
    def _generate_basic_imputation_suggestions(self, df: pd.DataFrame) -> str:
        """Generate basic imputation suggestions without AI"""
        suggestions = []
        missing_info = df.isnull().sum()
        cols_with_missing = missing_info[missing_info > 0]
        
        if cols_with_missing.empty:
            return "No missing values found in the dataset."
        
        suggestions.append("Basic imputation suggestions:")
        
        for col in cols_with_missing.index:
            if np.issubdtype(df[col].dtype, np.number):
                # For numeric columns
                suggestions.append(f"\nColumn: {col}")
                suggestions.append("- Consider median if there are outliers")
                suggestions.append("- Consider mean if the distribution is roughly symmetric")
                suggestions.append("- Consider forward fill if the data has a temporal component")
            else:
                # For categorical columns
                suggestions.append(f"\nColumn: {col}")
                suggestions.append("- Consider mode (most frequent value)")
                suggestions.append("- Consider a constant value like 'Unknown' or 'Missing'")
                suggestions.append("- Consider forward fill if the data has a logical sequence")
        
        return "\n".join(suggestions)

    def analyze_visualization(self, df: pd.DataFrame, viz_type: str, columns: List[str], plot_data: dict = None) -> str:
        """Analyze patterns and insights from visualizations"""
        # Prepare data statistics for the columns
        stats_info = []
        for col in columns:
            if col in df.columns:
                col_data = df[col].dropna()
                stats = {
                    'name': col,
                    'dtype': df[col].dtype,
                    'mean': col_data.mean() if np.issubdtype(df[col].dtype, np.number) else None,
                    'median': col_data.median() if np.issubdtype(df[col].dtype, np.number) else None,
                    'std': col_data.std() if np.issubdtype(df[col].dtype, np.number) else None,
                    'unique_count': col_data.nunique(),
                    'missing_count': df[col].isnull().sum()
                }
                stats_info.append(stats)
        
        # Format statistics with proper handling of None values
        def format_stat(value):
            return f"{value:.2f}" if value is not None else "N/A"
        
        # Construct prompt based on visualization type
        if viz_type == "histogram":
            prompt = f"""
            Analyze this histogram visualization for column '{columns[0]}':
            
            Column Statistics:
            - Data type: {stats_info[0]['dtype']}
            - Mean: {format_stat(stats_info[0]['mean'])}
            - Median: {format_stat(stats_info[0]['median'])}
            - Standard deviation: {format_stat(stats_info[0]['std'])}
            - Unique values: {stats_info[0]['unique_count']}
            
            Please analyze:
            1. The distribution shape (normal, skewed, bimodal, etc.)
            2. Any notable peaks or patterns
            3. Presence of potential outliers
            4. Practical implications of this distribution
            
            Keep the analysis concise and data-driven.
            """
            
        elif viz_type == "scatter":
            correlation = df[columns].corr().iloc[0, 1] if len(columns) == 2 else None
            prompt = f"""
            Analyze this scatter plot between '{columns[0]}' and '{columns[1]}':
            
            Correlation coefficient: {format_stat(correlation)}
            
            Variables:
            1. {columns[0]}:
               - Mean: {format_stat(stats_info[0]['mean'])}
               - Std: {format_stat(stats_info[0]['std'])}
            
            2. {columns[1]}:
               - Mean: {format_stat(stats_info[1]['mean'])}
               - Std: {format_stat(stats_info[1]['std'])}
            
            Please analyze:
            1. The relationship pattern (linear, non-linear, no relationship)
            2. Strength and direction of the relationship
            3. Any notable clusters or outliers
            4. Practical implications of this relationship
            
            Keep the analysis concise and data-driven.
            """
            
        elif viz_type == "boxplot":
            prompt = f"""
            Analyze this box plot for the following columns: {', '.join(columns)}
            
            Statistics for each column:
            """ + "\n".join([f"""
            {stats['name']}:
            - Median: {format_stat(stats['median'])}
            - Std: {format_stat(stats['std'])}
            """ for stats in stats_info]) + """
            
            Please analyze:
            1. The distribution of values and central tendencies
            2. Presence and severity of outliers
            3. Comparison between different columns (if multiple)
            4. Any notable patterns or anomalies
            
            Keep the analysis concise and data-driven.
            """
            
        elif viz_type == "line":
            prompt = f"""
            Analyze this line plot between '{columns[0]}' and '{columns[1]}':
            
            Variables:
            1. {columns[0]} (x-axis):
               - Unique values: {stats_info[0]['unique_count']}
            
            2. {columns[1]} (y-axis):
               - Mean: {format_stat(stats_info[1]['mean'])}
               - Std: {format_stat(stats_info[1]['std'])}
            
            Please analyze:
            1. Overall trend (increasing, decreasing, cyclical)
            2. Any notable patterns or seasonality
            3. Presence of significant peaks or troughs
            4. Any apparent anomalies or outliers
            
            Keep the analysis concise and data-driven.
            """
            
        elif viz_type == "heatmap":
            correlations = df[columns].corr()
            prompt = f"""
            Analyze this correlation heatmap for the following columns: {', '.join(columns)}
            
            Correlation matrix:
            {correlations.to_string()}
            
            Please analyze:
            1. The strongest positive and negative correlations
            2. Any notable correlation patterns or clusters
            3. Variables that show little to no correlation
            4. Practical implications of these relationships
            
            Keep the analysis concise and data-driven.
            """
        
        else:
            prompt = f"""
            Analyze this {viz_type} visualization for columns: {', '.join(columns)}
            
            Statistics for each column:
            """ + "\n".join([f"""
            {stats['name']}:
            - Data type: {stats['dtype']}
            - Mean: {format_stat(stats['mean'])}
            - Median: {format_stat(stats['median'])}
            - Std: {format_stat(stats['std'])}
            - Unique values: {stats['unique_count']}
            """ for stats in stats_info]) + """
            
            Please provide:
            1. Key patterns or trends in the visualization
            2. Notable outliers or anomalies
            3. Potential insights for business or analysis
            
            Keep the analysis concise and data-driven.
            """
        
        try:
            response = self.model.generate_content(prompt)
            return response.text
        except Exception as e:
            return f"Error analyzing visualization: {str(e)}"

class AnalyticsBot:
    def __init__(self, df: pd.DataFrame, gemini_api_key: str = None):
        self.df = df
        self.numeric_cols = self.df.select_dtypes(include=[np.number]).columns
        self.categorical_cols = self.df.select_dtypes(include=['object']).columns
        self.missing_handler = MissingValueHandler(df)
        self.insights_generator = InsightsGenerator(gemini_api_key) if gemini_api_key else InsightsGenerator("dummy_key")
    
    def process_query(self, query: str) -> Tuple[str, any]:
        query = query.lower()
        
        # Handle visualization suggestions
        if any(word in query for word in ['suggest', 'recommend']) and any(word in query for word in ['plot', 'graph', 'chart', 'viz', 'visual']):
            try:
                suggestions = self.insights_generator.suggest_visualizations(self.df)
                return "Here are suggested visualizations for your data:", suggestions
            except Exception as e:
                return f"Error generating visualization suggestions: {str(e)}", None
        
        # Handle imputation suggestions
        if ('suggest' in query or 'recommend' in query) and ('missing' in query or 'imputation' in query):
            try:
                suggestions = self.insights_generator.suggest_imputation_strategy(self.df)
                return "Here are suggested strategies for handling missing values:", suggestions
            except Exception as e:
                return f"Error generating imputation suggestions: {str(e)}", None
        
        # Handle insights requests
        if 'insight' in query or 'analyze' in query or 'tell me about' in query:
            try:
                insights = self.insights_generator.generate_insights(self.df)
                return "Here are the insights about your data:", insights
            except Exception as e:
                return f"Error generating insights: {str(e)}", None
        
        # Missing value handling queries
        if 'missing' in query:
            if 'show' in query or 'display' in query:
                missing_info = self.missing_handler.get_missing_info()
                return "Here's the missing value information:", missing_info
            
            # Handle imputation requests
            if 'impute' in query or 'fill' in query:
                method = self._extract_imputation_method(query)
                columns = self._extract_column_names(query)
                
                if method:
                    imputed_df, message = self.missing_handler.impute_values(method, columns)
                    self.df = imputed_df  # Update the dataframe with imputed values
                    return f"{message}. Here's the first few rows:", self.df.head()
                return "Please specify an imputation method (mean, median, mode, constant, ffill, or bfill)", None
        
        # Handle outlier detection requests
        elif 'outlier' in query:
            if 'zscore' in query or 'z-score' in query:
                method = 'zscore'
            else:
                method = 'iqr'
            
            # Determine available dimensions based on numeric columns
            dimensions = len(self.numeric_cols)
            dimension_options = ["1D"]
            if dimensions >= 2:
                dimension_options.append("2D")
                
            return ("Select analysis type and columns:", {
                "type": "outliers",
                "columns": list(self.numeric_cols),
                "method": method,
                "dimension_options": dimension_options
            })
        
        # Visualization queries
        elif any(plot_type in query for plot_type in ['plot', 'graph', 'chart', 'viz', 'visual']):
            if 'histogram' in query:
                return ("Please select a column for the histogram:", {
                    "type": "histogram",
                    "columns": list(self.numeric_cols)
                })
                
            elif 'box' in query or 'boxplot' in query:
                return ("Please select columns for the box plot:", {
                    "type": "boxplot",
                    "columns": list(self.numeric_cols)
                })
                
            elif 'scatter' in query:
                return ("Please select columns for the scatter plot:", {
                    "type": "scatter",
                    "numeric_columns": list(self.numeric_cols),
                    "categorical_columns": list(self.categorical_cols)
                })
                
            elif 'line' in query:
                return ("Please select columns for the line plot:", {
                    "type": "line",
                    "numeric_columns": list(self.numeric_cols),
                    "categorical_columns": list(self.categorical_cols)
                })
                
            elif 'bar' in query:
                return ("Please select columns for the bar plot:", {
                    "type": "bar",
                    "numeric_columns": list(self.numeric_cols),
                    "categorical_columns": list(self.categorical_cols)
                })
                
            elif 'correlation' in query or 'heatmap' in query:
                if len(self.numeric_cols) < 2:
                    return "Not enough numeric columns for a correlation heatmap.", None
                return ("Please select columns for the correlation heatmap:", {
                    "type": "heatmap",
                    "columns": list(self.numeric_cols)
                })
                
            elif 'violin' in query:
                return ("Please select columns for the violin plot:", {
                    "type": "violin",
                    "numeric_columns": list(self.numeric_cols),
                    "categorical_columns": list(self.categorical_cols)
                })
            
            else:
                return ("Please select the type of plot:", {
                    "type": "plot_selection",
                    "options": ["histogram", "boxplot", "scatter", "line", "bar", "heatmap", "violin"]
                })
        
        elif 'describe' in query or 'summary' in query:
            col = self._extract_column_name(query)
            if col in self.df.columns:
                stats = self.df[col].describe()
                return f"Here's the summary statistics for {col}:", stats
            return "Here's the summary statistics for all numeric columns:", self.df.describe()
            
        # Handle normalization requests
        elif any(word in query for word in ['normalize', 'normalise', 'scaling', 'scale']):
            method = None
            if 'minmax' in query or 'min-max' in query:
                method = 'minmax'
            elif 'zscore' in query or 'z-score' in query:
                method = 'zscore'
            elif 'robust' in query:
                method = 'robust'
            
            return ("Select normalization options:", {
                "type": "normalize",
                "columns": list(self.numeric_cols),
                "method": method,
                "methods": ["minmax", "zscore", "robust"]
            })
        
        # Handle duplicate detection/removal requests
        elif any(word in query for word in ['duplicate', 'duplicates', 'unique']):
            return ("Select options for handling duplicates:", {
                "type": "duplicates",
                "columns": list(self.df.columns)
            })
        
        # Default response
        return "I can help you with:\n- Mean/average of columns\n- Median of columns\n- Histograms\n- Correlation heatmap\n- Summary statistics\n\nTry asking something like 'Show me the histogram of Sales'", None

    def _extract_imputation_method(self, query: str) -> str:
        """Extract imputation method from query"""
        methods = {
            'mean': ['mean', 'average'],
            'median': ['median', 'middle'],
            'mode': ['mode', 'most frequent', 'most common'],
            'constant': ['constant', 'fixed', 'zero'],
            'ffill': ['forward fill', 'ffill', 'forward'],
            'bfill': ['backward fill', 'bfill', 'backward']
        }
        
        for method, keywords in methods.items():
            if any(keyword in query for keyword in keywords):
                return method
        return None
    
    def _extract_column_names(self, query: str) -> List[str]:
        """Extract multiple column names from query"""
        return [col for col in self.df.columns if col.lower() in query.lower()]

    def _extract_column_name(self, query: str) -> str:
        """Extract column name from query"""
        for col in self.df.columns:
            if col.lower() in query.lower():
                return col
        return ""

    def _create_histogram(self, column: str) -> Tuple[plt.Figure, str]:
        """Create histogram and analyze it"""
        fig, ax = plt.subplots()
        sns.histplot(data=self.df, x=column, ax=ax)
        plt.title(f'Histogram of {column}')
        analysis = self.insights_generator.analyze_visualization(self.df, "histogram", [column])
        return fig, analysis

    def _create_scatter_plot(self, x_col: str, y_col: str, hue_col: str = None) -> Tuple[plt.Figure, str]:
        """Create scatter plot and analyze it"""
        fig, ax = plt.subplots(figsize=(10, 6))
        sns.scatterplot(data=self.df, x=x_col, y=y_col, hue=hue_col, ax=ax)
        plt.title(f'Scatter Plot: {x_col} vs {y_col}')
        analysis = self.insights_generator.analyze_visualization(self.df, "scatter", [x_col, y_col])
        return fig, analysis

    def _create_boxplot(self, columns: List[str]) -> Tuple[plt.Figure, str]:
        """Create box plot and analyze it"""
        fig, ax = plt.subplots(figsize=(10, 6))
        sns.boxplot(data=self.df[columns], ax=ax)
        plt.xticks(rotation=45)
        plt.title('Box Plot')
        analysis = self.insights_generator.analyze_visualization(self.df, "boxplot", columns)
        return fig, analysis

    def _create_line_plot(self, x_col: str, y_col: str, hue_col: str = None) -> Tuple[plt.Figure, str]:
        """Create line plot and analyze it"""
        fig, ax = plt.subplots(figsize=(10, 6))
        sns.lineplot(data=self.df, x=x_col, y=y_col, hue=hue_col, ax=ax)
        plt.xticks(rotation=45)
        plt.title(f'Line Plot: {x_col} vs {y_col}')
        analysis = self.insights_generator.analyze_visualization(self.df, "line", [x_col, y_col])
        return fig, analysis

    def _create_bar_plot(self, x_col: str, y_col: str, hue_col: str = None) -> Tuple[plt.Figure, str]:
        """Create bar plot and analyze it"""
        fig, ax = plt.subplots(figsize=(10, 6))
        sns.barplot(data=self.df, x=x_col, y=y_col, hue=hue_col, ax=ax)
        plt.xticks(rotation=45)
        plt.title(f'Bar Plot: {x_col} vs {y_col}')
        analysis = self.insights_generator.analyze_visualization(self.df, "bar", [x_col, y_col])
        return fig, analysis

    def _create_violin_plot(self, x_col: str, y_col: str) -> Tuple[plt.Figure, str]:
        """Create violin plot and analyze it"""
        fig, ax = plt.subplots(figsize=(10, 6))
        sns.violinplot(data=self.df, x=x_col, y=y_col, ax=ax)
        plt.xticks(rotation=45)
        plt.title(f'Violin Plot: {x_col} vs {y_col}')
        analysis = self.insights_generator.analyze_visualization(self.df, "violin", [x_col, y_col])
        return fig, analysis

    def _create_correlation_heatmap(self, columns: List[str]) -> plt.Figure:
        """Create correlation heatmap"""
        fig, ax = plt.subplots(figsize=(10, 8))
        sns.heatmap(self.df[columns].corr(), annot=True, ax=ax)
        plt.title('Correlation Heatmap')
        return fig

    def _detect_outliers(self, column: str, method: str = 'iqr', y_column: str = None) -> Tuple[pd.Series, plt.Figure]:
        """
        Detect outliers in a specified column using either IQR or Z-score method
        
        Args:
            column: Name of the column to analyze
            method: 'iqr' or 'zscore'
            y_column: Name of the second column for 2D analysis
            
        Returns:
            Tuple of (outlier_info, visualization)
        """
        if column not in self.numeric_cols:
            return f"Column {column} is not numeric. Outlier detection requires numeric data.", None
            
        data = self.df[column].dropna()
        outliers = pd.Series()
        
        if method == 'iqr':
            Q1 = data.quantile(0.25)
            Q3 = data.quantile(0.75)
            IQR = Q3 - Q1
            lower_bound = Q1 - 1.5 * IQR
            upper_bound = Q3 + 1.5 * IQR
            outliers = data[(data < lower_bound) | (data > upper_bound)]
            
            # Create visualization
            fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(15, 5))
            
            # Box plot
            sns.boxplot(data=data, ax=ax1)
            ax1.set_title(f'Box Plot of {column}')
            
            # Scatter plot highlighting outliers
            sns.scatterplot(data=self.df.reset_index(), x='index', y=column, ax=ax2)
            if len(outliers) > 0:
                outlier_indices = outliers.index
                sns.scatterplot(x=outlier_indices, y=outliers.values, color='red', ax=ax2)
            ax2.set_title(f'Outliers in {column} (red points)')
            plt.tight_layout()
            
        elif method == 'zscore':
            z_scores = np.abs((data - data.mean()) / data.std())
            outliers = data[z_scores > 3]  # Points beyond 3 standard deviations
            
            # Create visualization
            fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(15, 5))
            
            # Distribution plot with z-score boundaries
            sns.histplot(data=data, ax=ax1)
            ax1.axvline(data.mean() + 3*data.std(), color='r', linestyle='--')
            ax1.axvline(data.mean() - 3*data.std(), color='r', linestyle='--')
            ax1.set_title(f'Distribution of {column} with Z-score Boundaries')
            
            # Scatter plot highlighting outliers
            sns.scatterplot(data=self.df.reset_index(), x='index', y=column, ax=ax2)
            if len(outliers) > 0:
                outlier_indices = outliers.index
                sns.scatterplot(x=outlier_indices, y=outliers.values, color='red', ax=ax2)
            ax2.set_title(f'Outliers in {column} (red points)')
            plt.tight_layout()
        
        # Prepare outlier information
        outlier_info = pd.DataFrame({
            'Index': outliers.index,
            'Value': outliers.values,
            'Method': method
        })
        
        return outlier_info, fig

    def _normalize_data(self, method: str, columns: List[str] = None) -> Tuple[pd.DataFrame, str, plt.Figure]:
        """
        Normalize data using specified method
        Methods: minmax, zscore, robust
        """
        df_normalized = self.df.copy()
        
        if columns is None:
            columns = list(self.numeric_cols)
        else:
            columns = [col for col in columns if col in self.numeric_cols]
            
        if not columns:
            return df_normalized, "No numeric columns to normalize", None
        
        try:
            if method == 'minmax':
                # Min-Max scaling: (x - min) / (max - min)
                for col in columns:
                    min_val = df_normalized[col].min()
                    max_val = df_normalized[col].max()
                    df_normalized[col] = (df_normalized[col] - min_val) / (max_val - min_val)
                message = "Data normalized using Min-Max scaling (range: 0 to 1)"
                
            elif method == 'zscore':
                # Z-score scaling: (x - mean) / std
                for col in columns:
                    mean_val = df_normalized[col].mean()
                    std_val = df_normalized[col].std()
                    df_normalized[col] = (df_normalized[col] - mean_val) / std_val
                message = "Data normalized using Z-score scaling (mean=0, std=1)"
                
            elif method == 'robust':
                # Robust scaling: (x - median) / IQR
                for col in columns:
                    median_val = df_normalized[col].median()
                    q1 = df_normalized[col].quantile(0.25)
                    q3 = df_normalized[col].quantile(0.75)
                    iqr = q3 - q1
                    df_normalized[col] = (df_normalized[col] - median_val) / iqr
                message = "Data normalized using Robust scaling (median=0, IQR=1)"
                
            # Create comparison visualizations
            fig, axes = plt.subplots(len(columns), 2, figsize=(12, 4*len(columns)))
            if len(columns) == 1:
                axes = axes.reshape(1, -1)
            
            for idx, col in enumerate(columns):
                # Original distribution
                sns.histplot(data=self.df[col], ax=axes[idx, 0])
                axes[idx, 0].set_title(f'Original Distribution - {col}')
                
                # Normalized distribution
                sns.histplot(data=df_normalized[col], ax=axes[idx, 1])
                axes[idx, 1].set_title(f'Normalized Distribution - {col}')
            
            plt.tight_layout()
            
            return df_normalized, message, fig
            
        except Exception as e:
            return df_normalized, f"Error during normalization: {str(e)}", None

    def _handle_duplicates(self, columns: List[str] = None) -> Tuple[pd.DataFrame, Dict, plt.Figure]:
        """
        Detect and handle duplicate rows in the dataset
        
        Args:
            columns: Specific columns to check for duplicates. If None, checks all columns
            
        Returns:
            Tuple of (DataFrame with duplicates removed, duplicate info, visualization)
        """
        # Create copy of dataframe
        df_clean = self.df.copy()
        
        # Get duplicate information
        if columns:
            duplicates = df_clean[df_clean.duplicated(subset=columns, keep='first')]
            duplicate_mask = df_clean.duplicated(subset=columns, keep=False)
        else:
            duplicates = df_clean[df_clean.duplicated(keep='first')]
            duplicate_mask = df_clean.duplicated(keep=False)
        
        # Prepare duplicate information
        duplicate_info = {
            'total_duplicates': len(duplicates),
            'duplicate_rows': duplicates.index.tolist(),
            'columns_checked': columns if columns else 'all columns',
            'unique_rows': len(df_clean[~duplicate_mask])
        }
        
        # Create visualization
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(15, 5))
        
        # Pie chart of duplicate vs unique rows
        sizes = [duplicate_info['unique_rows'], duplicate_info['total_duplicates']]
        labels = ['Unique Rows', 'Duplicate Rows']
        ax1.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
        ax1.set_title('Distribution of Duplicate Rows')
        
        # Bar chart of duplicate counts by column (if specific columns selected)
        if columns:
            duplicate_counts = []
            for col in columns:
                dup_count = df_clean[df_clean.duplicated(subset=[col], keep=False)].shape[0]
                duplicate_counts.append(dup_count)
            
            ax2.bar(columns, duplicate_counts)
            ax2.set_title('Duplicate Counts by Column')
            ax2.set_xticklabels(columns, rotation=45)
            ax2.set_ylabel('Number of Duplicates')
        else:
            ax2.text(0.5, 0.5, 'Checking all columns for duplicates', 
                    horizontalalignment='center', verticalalignment='center')
            ax2.axis('off')
        
        plt.tight_layout()
        
        return df_clean[~duplicate_mask], duplicate_info, fig

def initialize_session_state():
    """Initialize session state variables"""
    if 'messages' not in st.session_state:
        st.session_state.messages = []
    if 'df' not in st.session_state:
        st.session_state.df = None
    if 'bot' not in st.session_state:
        st.session_state.bot = None

def generate_pdf_report(messages) -> bytes:
    """Generate a structured PDF report from chat messages"""
    pdf = FPDF()
    pdf.add_page()
    
    # Title and Header
    pdf.set_font("Arial", "B", 20)
    pdf.cell(0, 15, "Data Analysis Report", ln=True, align='C')
    pdf.set_font("Arial", "I", 12)
    pdf.cell(0, 10, f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True, align='C')
    pdf.ln(10)
    
    # Dataset Overview Section
    if hasattr(st.session_state.bot, 'df'):
        df = st.session_state.bot.df
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, "1. Dataset Overview", ln=True)
        pdf.ln(5)
        
        pdf.set_font("Arial", "B", 12)
        pdf.cell(0, 10, "Basic Information:", ln=True)
        pdf.set_font("Arial", "", 12)
        pdf.multi_cell(0, 8, f"- Number of Records: {len(df)}")
        pdf.multi_cell(0, 8, f"- Number of Features: {len(df.columns)}")
        pdf.multi_cell(0, 8, f"- Data Types: {', '.join(df.dtypes.astype(str).unique())}")
        pdf.ln(5)
        
        # Missing Values Summary
        pdf.set_font("Arial", "B", 12)
        pdf.cell(0, 10, "Missing Values Summary:", ln=True)
        pdf.set_font("Arial", "", 12)
        missing_info = st.session_state.bot.missing_handler.get_missing_info()
        has_missing = False
        for col in missing_info.index:
            total_missing = missing_info.loc[col, 'Total Missing']
            if total_missing > 0:
                has_missing = True
                percent_missing = missing_info.loc[col, 'Percent Missing']
                pdf.multi_cell(0, 8, f"- {col}: {total_missing} values ({percent_missing}%)")
        if not has_missing:
            pdf.multi_cell(0, 8, "No missing values found in the dataset.")
        pdf.ln(10)
    
    # Analysis Section
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, "2. Analysis Results", ln=True)
    pdf.ln(5)
    
    # Process each message
    analysis_count = 1
    for msg in messages:
        if msg["role"] == "user":
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 10, f"Query {analysis_count}:", ln=True)
            pdf.set_font("Arial", "", 12)
            pdf.multi_cell(0, 8, msg["content"])
            pdf.ln(5)
            
        else:  # assistant's response
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 10, f"Analysis {analysis_count}:", ln=True)
            pdf.set_font("Arial", "", 12)
            
            # Handle different types of responses
            if isinstance(msg["content"], tuple):
                message, content = msg["content"]
                
                # Add the message
                pdf.multi_cell(0, 8, str(message))
                
                # Handle different types of content
                if isinstance(content, pd.DataFrame):
                    pdf.ln(5)
                    pdf.set_font("Arial", "B", 12)
                    pdf.cell(0, 10, "Results:", ln=True)
                    pdf.set_font("Arial", "", 12)
                    pdf.multi_cell(0, 8, content.to_string())
                    
                elif isinstance(content, str):
                    pdf.multi_cell(0, 8, content)
                    
                elif isinstance(content, dict):
                    if "type" in content:
                        viz_type = content["type"]
                        if viz_type in ["histogram", "scatter", "boxplot", "line", "bar", "heatmap", "violin"]:
                            pdf.ln(5)
                            pdf.set_font("Arial", "B", 12)
                            pdf.cell(0, 10, f"{viz_type.title()} Visualization:", ln=True)
                            pdf.set_font("Arial", "", 12)
                            
                            # Save the current figure if it exists
                            if plt.get_fignums():
                                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                                plt.savefig(temp_file.name, bbox_inches='tight', dpi=300)
                                plt.close()
                                pdf.image(temp_file.name, x=10, w=190)
                                temp_file.close()
                                os.unlink(temp_file.name)
                            
                            # Add any analysis text
                            if "analysis" in content:
                                pdf.ln(5)
                                pdf.set_font("Arial", "B", 12)
                                pdf.cell(0, 10, "Visual Analysis:", ln=True)
                                pdf.set_font("Arial", "", 12)
                                pdf.multi_cell(0, 8, str(content["analysis"]))
                    else:
                        for key, value in content.items():
                            if key not in ["type", "columns", "methods"]:
                                pdf.multi_cell(0, 8, f"{key}: {value}")
                
                elif isinstance(content, plt.Figure):
                    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                    content.savefig(temp_file.name, bbox_inches='tight', dpi=300)
                    plt.close(content)
                    pdf.image(temp_file.name, x=10, w=190)
                    temp_file.close()
                    os.unlink(temp_file.name)
            
            else:
                pdf.multi_cell(0, 8, str(msg["content"]))
            
            pdf.ln(10)
            analysis_count += 1
    
    # Statistical Summary Section
    if hasattr(st.session_state.bot, 'df'):
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, "3. Statistical Summary", ln=True)
        pdf.ln(5)
        
        # Numeric columns summary
        numeric_cols = df.select_dtypes(include=[np.number])
        if not numeric_cols.empty:
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 10, "Numerical Features:", ln=True)
            pdf.set_font("Arial", "", 12)
            summary_stats = numeric_cols.describe().round(2)
            pdf.multi_cell(0, 8, summary_stats.to_string())
            pdf.ln(5)
        
        # Categorical columns summary
        categorical_cols = df.select_dtypes(exclude=[np.number])
        if not categorical_cols.empty:
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 10, "Categorical Features:", ln=True)
            pdf.set_font("Arial", "", 12)
            for col in categorical_cols.columns:
                value_counts = df[col].value_counts().head(5)
                pdf.multi_cell(0, 8, f"\n{col} (top 5 categories):")
                for val, count in value_counts.items():
                    pdf.multi_cell(0, 8, f"- {val}: {count}")
    
    return pdf.output(dest='S').encode('latin-1')

def generate_word_report(messages) -> bytes:
    """Generate a Word document report from chat messages"""
    doc = Document()
    
    # Add title
    doc.add_heading('Data Analysis Report', 0)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    # Process each message
    for msg in messages:
        if msg["role"] == "user":
            doc.add_heading('User Query:', level=1)
            doc.add_paragraph(msg["content"])
            
        else:  # assistant's response
            doc.add_heading('Analysis Results:', level=1)
            
            # Handle different types of responses
            if isinstance(msg["content"], tuple):
                message, content = msg["content"]
                
                # Add the message
                doc.add_paragraph(str(message))
                
                # Handle different types of content
                if isinstance(content, pd.DataFrame):
                    # Convert DataFrame to table
                    t = doc.add_table(content.shape[0]+1, content.shape[1])
                    # Add headers
                    for j in range(content.shape[1]):
                        t.cell(0,j).text = str(content.columns[j])
                    # Add data
                    for i in range(content.shape[0]):
                        for j in range(content.shape[1]):
                            t.cell(i+1,j).text = str(content.iloc[i,j])
                    
                elif isinstance(content, str):
                    doc.add_paragraph(content)
                    
                elif isinstance(content, dict):
                    if "type" in content:
                        viz_type = content["type"]
                        if viz_type in ["histogram", "scatter", "boxplot", "line", "bar", "heatmap", "violin"]:
                            # Save the current figure if it exists
                            if plt.get_fignums():
                                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                                plt.savefig(temp_file.name, bbox_inches='tight', dpi=300)
                                plt.close()
                                # Add the image to document
                                doc.add_picture(temp_file.name, width=Inches(6))
                                temp_file.close()
                                os.unlink(temp_file.name)
                            
                            # Add any analysis text
                            if "analysis" in content:
                                doc.add_paragraph(str(content["analysis"]))
                    else:
                        for key, value in content.items():
                            if key not in ["type", "columns", "methods"]:
                                doc.add_paragraph(f"{key}: {value}")
                
                elif isinstance(content, plt.Figure):
                    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                    content.savefig(temp_file.name, bbox_inches='tight', dpi=300)
                    plt.close(content)
                    doc.add_picture(temp_file.name, width=Inches(6))
                    temp_file.close()
                    os.unlink(temp_file.name)
            
            else:
                doc.add_paragraph(str(msg["content"]))
    
    # Add dataset summary if available
    if hasattr(st.session_state.bot, 'df'):
        doc.add_heading('Dataset Summary', level=1)
        df = st.session_state.bot.df
        doc.add_paragraph(f"Number of rows: {len(df)}")
        doc.add_paragraph(f"Number of columns: {len(df.columns)}")
        
        # Add summary statistics
        doc.add_heading('Summary Statistics', level=2)
        summary_stats = df.describe()
        t = doc.add_table(summary_stats.shape[0]+1, summary_stats.shape[1])
        # Add headers
        for j in range(summary_stats.shape[1]):
            t.cell(0,j).text = str(summary_stats.columns[j])
        # Add data
        for i in range(summary_stats.shape[0]):
            for j in range(summary_stats.shape[1]):
                t.cell(i+1,j).text = str(summary_stats.iloc[i,j])
    
    # Save to bytes buffer
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer.getvalue()

def main():
    st.title("Analytics Chatbot")
    
    # Initialize session state if needed
    if "bot" not in st.session_state or st.session_state.bot is None:
        st.session_state.bot = None
        st.session_state.messages = []
    
    # Add download options in sidebar for chat history
    if st.session_state.messages:  # Check if there are any messages
        report_format = st.sidebar.selectbox(
            "Select report format:",
            ["PDF", "DOCX"]
        )
        
        if report_format == "PDF":
            pdf_bytes = generate_pdf_report(st.session_state.messages)
            st.sidebar.download_button(
                label="Download Analysis Report (PDF)",
                data=pdf_bytes,
                file_name=f"analysis_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                mime="application/pdf"
            )
        else:
            docx_buffer = generate_word_report(st.session_state.messages)
            st.sidebar.download_button(
                label="Download Analysis Report (DOCX)",
                data=docx_buffer,
                file_name=f"analysis_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        
        # Option to include visualizations
        if st.sidebar.checkbox("Include visualizations"):
            st.sidebar.info("Visualizations will be included in the report")
    
    # Add dataset download options in sidebar if data exists and has been modified
    if (st.session_state.bot is not None and 
        hasattr(st.session_state.bot, 'df') and 
        st.session_state.bot.df is not None):
        st.sidebar.markdown("### Download Modified Dataset")
        file_format = st.sidebar.selectbox(
            "Select file format:",
            ["CSV", "Excel", "JSON"]
        )
        
        if file_format == "CSV":
            csv_data = st.session_state.bot.df.to_csv(index=False)
            st.sidebar.download_button(
                label="Download Dataset as CSV",
                data=csv_data,
                file_name=f"modified_dataset_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv"
            )
        elif file_format == "Excel":
            excel_buffer = BytesIO()
            st.session_state.bot.df.to_excel(excel_buffer, index=False)
            excel_buffer.seek(0)
            st.sidebar.download_button(
                label="Download Dataset as Excel",
                data=excel_buffer,
                file_name=f"modified_dataset_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
        else:  # JSON
            json_data = st.session_state.bot.df.to_json()
            st.sidebar.download_button(
                label="Download Dataset as JSON",
                data=json_data,
                file_name=f"modified_dataset_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                mime="application/json"
            )
    
    initialize_session_state()

    # File upload
    uploaded_file = st.file_uploader("Choose a CSV file", type="csv")
    
    if uploaded_file is not None:
        # Read the data and initialize the bot
        if st.session_state.df is None:
            st.session_state.df = pd.read_csv(uploaded_file)
            st.session_state.bot = AnalyticsBot(st.session_state.df)
            st.session_state.messages.append({
                "role": "assistant",
                "content": "I've loaded your data! You can ask me questions like:\n" +
                          "- Generate insights about the data\n" +
                          "- Generate Gemini insights\n" +
                          "- Tell me about this dataset\n" +
                          "- Analyze the sales trends\n" +
                          "- Show me missing values\n" +
                          "- Impute missing values using mean/median/mode\n" +
                          "- Fill missing values with forward fill\n" +
                          "- What's the mean of [column]?\n" +
                          "- Show me a histogram of [column]\n" +
                          "- Show me the correlation heatmap\n" +
                          "- Describe [column]"
            })

        # Display chat messages
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                if isinstance(message["content"], tuple):
                    st.write(message["content"][0])
                    if message["content"][1] is not None:
                        if isinstance(message["content"][1], dict):
                            # Handle visualization requests
                            if message["content"][1]["type"] == "histogram":
                                selected_col = st.selectbox(
                                    "Select a column:",
                                    message["content"][1]["columns"]
                                )
                                if selected_col:
                                    fig, analysis = st.session_state.bot._create_histogram(selected_col)
                                    st.pyplot(fig)
                                    plt.close()
                                    if analysis:
                                        st.write("Analysis:")
                                        st.write(analysis)
                            
                            elif message["content"][1]["type"] == "boxplot":
                                selected_cols = st.multiselect(
                                    "Select columns:",
                                    message["content"][1]["columns"],
                                    default=[message["content"][1]["columns"][0]] if len(message["content"][1]["columns"]) > 0 else []
                                )
                                if selected_cols:
                                    fig, analysis = st.session_state.bot._create_boxplot(selected_cols)
                                    st.pyplot(fig)
                                    plt.close()
                                    if analysis:
                                        st.write("Analysis:")
                                        st.write(analysis)
                            
                            elif message["content"][1]["type"] in ["scatter", "line", "bar"]:
                                x_col = st.selectbox(
                                    "Select X-axis column:",
                                    message["content"][1]["numeric_columns"] + message["content"][1]["categorical_columns"]
                                )
                                y_col = st.selectbox(
                                    "Select Y-axis column:",
                                    message["content"][1]["numeric_columns"]
                                )
                                hue_col = st.selectbox(
                                    "Select grouping column (optional):",
                                    ["None"] + list(message["content"][1]["categorical_columns"])
                                )
                                
                                if x_col and y_col:
                                    hue = None if hue_col == "None" else hue_col
                                    if message["content"][1]["type"] == "scatter":
                                        fig, analysis = st.session_state.bot._create_scatter_plot(x_col, y_col, hue)
                                    elif message["content"][1]["type"] == "line":
                                        fig, analysis = st.session_state.bot._create_line_plot(x_col, y_col, hue)
                                    else:  # bar plot
                                        fig, analysis = st.session_state.bot._create_bar_plot(x_col, y_col, hue)
                                    st.pyplot(fig)
                                    plt.close()
                                    if analysis:
                                        st.write("Analysis:")
                                        st.write(analysis)
                            
                            elif message["content"][1]["type"] == "violin":
                                x_col = st.selectbox(
                                    "Select X-axis (categorical) column:",
                                    message["content"][1]["categorical_columns"]
                                )
                                y_col = st.selectbox(
                                    "Select Y-axis (numeric) column:",
                                    message["content"][1]["numeric_columns"]
                                )
                                if x_col and y_col:
                                    fig, analysis = st.session_state.bot._create_violin_plot(x_col, y_col)
                                    st.pyplot(fig)
                                    plt.close()
                                    if analysis:
                                        st.write("Analysis:")
                                        st.write(analysis)
                            
                            elif message["content"][1]["type"] == "heatmap":
                                selected_cols = st.multiselect(
                                    "Select columns for correlation analysis:",
                                    message["content"][1]["columns"],
                                    default=list(message["content"][1]["columns"])[:min(5, len(message["content"][1]["columns"]))]
                                )
                                if selected_cols:
                                    fig = st.session_state.bot._create_correlation_heatmap(selected_cols)
                                    st.pyplot(fig)
                                    plt.close()
                                    if analysis:
                                        st.write("Analysis:")
                                        st.write(analysis)
                            
                            elif message["content"][1]["type"] == "plot_selection":
                                plot_type = st.selectbox(
                                    "Select the type of plot you want to create:",
                                    message["content"][1]["options"]
                                )
                                if plot_type:
                                    new_response = st.session_state.bot.process_query(f"show {plot_type}")
                                    st.session_state.messages.append({"role": "assistant", "content": new_response})
                                    st.rerun()
                            
                            elif message["content"][1]["type"] == "outliers":
                                # Let user select dimension first
                                dimension = st.selectbox(
                                    "Select analysis type:",
                                    message["content"][1]["dimension_options"]
                                )
                                
                                if dimension == "2D":
                                    col1, col2 = st.columns(2)
                                    with col1:
                                        x_col = st.selectbox(
                                            "Select X-axis column:",
                                            message["content"][1]["columns"]
                                        )
                                    with col2:
                                        y_col = st.selectbox(
                                            "Select Y-axis column:",
                                            [col for col in message["content"][1]["columns"] if col != x_col]
                                        )
                                    
                                    if x_col and y_col:
                                        outlier_info, fig = st.session_state.bot._detect_outliers(
                                            x_col,
                                            method=message["content"][1]["method"],
                                            y_column=y_col
                                        )
                                        
                                        if isinstance(outlier_info, pd.DataFrame):
                                            st.write(f"Found {len(outlier_info)} outliers")
                                            st.dataframe(outlier_info)
                                            st.pyplot(fig)
                                            plt.close()
                                        else:
                                            st.write(outlier_info)
                                else:  # 1D analysis
                                    selected_col = st.selectbox(
                                        "Select column for analysis:",
                                        message["content"][1]["columns"]
                                    )
                                    
                                    if selected_col:
                                        outlier_info, fig = st.session_state.bot._detect_outliers(
                                            selected_col, 
                                            method=message["content"][1]["method"]
                                        )
                                        
                                        if isinstance(outlier_info, pd.DataFrame):
                                            st.write(f"Found {len(outlier_info)} outliers in {selected_col}")
                                            st.dataframe(outlier_info)
                                            st.pyplot(fig)
                                            plt.close()
                                        else:
                                            st.write(outlier_info)
                                            
                                # Add method explanation
                                with st.expander("About the outlier detection method"):
                                    if message["content"][1]["method"] == "iqr":
                                        st.write("""
                                        **IQR (Interquartile Range) Method:**
                                        - Calculates Q1 (25th percentile) and Q3 (75th percentile)
                                        - IQR = Q3 - Q1
                                        - Points beyond Q1 - 1.5*IQR or Q3 + 1.5*IQR are considered outliers
                                        - More robust to extreme values than Z-score method
                                        """)
                                    else:
                                        st.write("""
                                        **Z-score Method:**
                                        - Calculates how many standard deviations a point is from the mean
                                        - Points beyond 3 standard deviations are considered outliers
                                        - Assumes data is normally distributed
                                        - More sensitive to extreme values than IQR method
                                        """)
                        elif isinstance(message["content"][1], plt.Figure):
                            st.pyplot(message["content"][1])
                            plt.close()
                        elif isinstance(message["content"][1], pd.DataFrame):
                            st.dataframe(message["content"][1])
                        else:
                            st.write(message["content"][1])
                else:
                    st.write(message["content"])

        # Chat input
        if prompt := st.chat_input("Ask me about your data!"):
            st.session_state.messages.append({"role": "user", "content": prompt})
            response = st.session_state.bot.process_query(prompt)
            st.session_state.messages.append({"role": "assistant", "content": response})
            st.rerun()

    else:
        st.write("Please upload a CSV file to begin analysis.")

if __name__ == "__main__":
    main()